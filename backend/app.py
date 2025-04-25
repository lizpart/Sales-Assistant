from fastapi import FastAPI, HTTPException, Depends, Query, Body
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from typing import List, Dict, Any, Optional
import requests
from requests.auth import HTTPBasicAuth
import os
from dotenv import load_dotenv
import json
from datetime import datetime, timedelta
import random
from langchain_google_genai import ChatGoogleGenerativeAI
import pandas as pd
from sklearn.model_selection import train_test_split
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import accuracy_score
import numpy as np
import io
from fastapi.responses import StreamingResponse

# Load environment variables
load_dotenv()

app = FastAPI(title="Sales Conversion Assistant API")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, replace with specific origin
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# CRM credentials
BC_ERP_URL = os.getenv("BC_ERP_URL", "https://bctest.dayliff.com:7048/BC160/ODataV4/Company('KENYA')/CRM_Interaction_Log_Entries")
BC_ERP_USERNAME = os.getenv("BC_ERP_USERNAME", "webservice")
BC_ERP_PASSWORD = os.getenv("BC_ERP_PASSWORD", "iqZwQDaYj665WV0aOgbSYFCDHsT9GxSxOUTTwOr5IV0=")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")

# Initialize LLM
llm = ChatGoogleGenerativeAI(
    model="gemini-1.5-pro",  # Using a powerful model for complex tasks
    temperature=0.7,
    google_api_key=GEMINI_API_KEY,
)

# Authentication for CRM requests
def get_auth():
    return HTTPBasicAuth(BC_ERP_USERNAME, BC_ERP_PASSWORD)

# Models
class PumpDetails(BaseModel):
    model_name: str = Field(description="The model name of the pump")

class Customer(BaseModel):
    id: str
    name: str
    contact_info: str
    industry: Optional[str] = None
    purchase_history: Optional[List[Dict[str, Any]]] = None

class SalesInteraction(BaseModel):
    customer_id: str
    interaction_type: str
    date: str
    notes: str
    objections: Optional[List[str]] = None
    next_steps: Optional[str] = None

class ProposalRequest(BaseModel):
    customer_name: str
    products: List[Dict[str, Any]]
    customer_requirements: str
    pain_points: Optional[List[str]] = None
    budget_range: Optional[str] = None

# Get pump details from ERP
def get_pump_from_erp(model_name: str):
    """
    Retrieves pump details for a given Davis & Shirtliff pump model from the Business Central ERP system.
    """
    try:
        # For products endpoint
        url = "https://bctest.dayliff.com:7048/BC160/ODataV4/Company('KENYA')/Items"
        
        params = {
            "$filter": f"contains(Description,'{model_name}') or contains(No,'{model_name}')"
        }
        
        response = requests.get(url, params=params, auth=get_auth())
        response.raise_for_status()

        pumps = response.json().get("value", [])

        # Find matching model
        for pump in pumps:
            pump_no = pump.get("No")
            pump_description = pump.get("Description", "").lower()
            
            if model_name.lower() in pump_description.lower() or model_name.lower() in pump_no.lower():
                return {
                    "Model_Number": pump_no,
                    "Description": pump_description,
                    "Inventory": pump.get("Inventory", "N/A"),
                    "Unit_Price": pump.get("Unit_Price", "N/A")
                }

        return {"error": f"No pump found with model name matching '{model_name}'. Please check the model name and try again."}

    except Exception as e:
        return {"error": str(e)}

# API Endpoints
@app.get("/")
def read_root():
    return {"message": "Welcome to Sales Conversion Assistant API"}

@app.get("/api/health")
def health_check():
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}

@app.post("/api/pump-details")
def get_pump_details(pump_details: PumpDetails):
    result = get_pump_from_erp(pump_details.model_name)
    if "error" in result:
        raise HTTPException(status_code=404, detail=result["error"])
    return result

@app.get("/api/sales-interactions")
def get_sales_interactions(customer_id: Optional[str] = None, limit: int = 10):
    """Get sales interactions from CRM"""
    try:
        url = BC_ERP_URL
        params = {}
        if customer_id:
            params["$filter"] = f"Contact_No eq '{customer_id}'"
        params["$top"] = str(limit)
        
        response = requests.get(url, params=params, auth=get_auth())
        response.raise_for_status()
        
        return response.json().get("value", [])
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/sales-invoices")
def get_sales_invoices(limit: int = 10):
    """Get sales invoices from CRM"""
    try:
        url = "https://bctest.dayliff.com:7048/BC160/ODataV4/Company('KENYA')/Sales_Invoice"
        params = {
            "$top": str(limit)
        }
        
        response = requests.get(url, params=params, auth=get_auth())
        response.raise_for_status()
        
        return response.json().get("value", [])
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/invoice-items/{invoice_id}")
def get_invoice_items(invoice_id: str):
    """Get items from a specific invoice"""
    try:
        url = "https://bctest.dayliff.com:7048/BC160/ODataV4/Company('KENYA')/Sales_InvoiceSalesLines"
        params = {
            "$filter": f"Document_No eq '{invoice_id}'"
        }
        
        response = requests.get(url, params=params, auth=get_auth())
        response.raise_for_status()
        
        return response.json().get("value", [])
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/analyze-objections")
def analyze_objections(interaction: SalesInteraction):
    """Analyze sales objections and provide recommendations"""
    if not interaction.objections:
        return {"message": "No objections provided for analysis"}
    
    objection_prompt = f"""
    Analyze the following sales objections and provide effective counter-strategies:
    Customer: {interaction.customer_id}
    Interaction Type: {interaction.interaction_type}
    Objections: {', '.join(interaction.objections)}
    Notes: {interaction.notes}
    
    Please provide:
    1. Analysis of each objection
    2. Recommended counter-strategies
    3. Suggested follow-up approach
    4. Key phrases to use in response
    """
    
    try:
        response = llm.invoke(objection_prompt)
        return {
            "analysis": response.content,
            "summary": "Analysis of objections and counter-strategies provided"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI analysis failed: {str(e)}")

@app.post("/api/generate-proposal")
def generate_proposal(request: ProposalRequest):
    """Generate a sales proposal document"""
    products_text = "\n".join([f"- {p['name']}: {p.get('description', 'N/A')}, Price: ${p.get('price', 'N/A')}" for p in request.products])
    pain_points = "\n".join([f"- {point}" for point in request.pain_points]) if request.pain_points else "No specific pain points mentioned"
    
    proposal_prompt = f"""
    Create a professional sales proposal document for:
    Customer: {request.customer_name}
    Budget Range: {request.budget_range or 'Not specified'}
    
    Customer Requirements:
    {request.customer_requirements}
    
    Pain Points to Address:
    {pain_points}
    
    Recommended Products/Solutions:
    {products_text}
    
    Create a complete, professional sales proposal document that includes:
    1. Executive Summary
    2. Customer Needs Assessment
    3. Proposed Solution with detailed product specifications
    4. Implementation Plan
    5. Pricing and ROI Analysis
    6. Next Steps
    
    Format this as a professional document with clear sections.
    """
    
    try:
        response = llm.invoke(proposal_prompt)
        
        # Create a Word-like document (for demonstration)
        from docx import Document
        doc = Document()
        
        # Add title
        doc.add_heading(f'Sales Proposal for {request.customer_name}', 0)
        
        # Add content (split by headings)
        content = response.content
        sections = content.split('\n# ')
        
        for section in sections:
            if section.strip():
                lines = section.split('\n', 1)
                if len(lines) > 1:
                    heading, content = lines
                    doc.add_heading(heading.strip('# '), level=1)
                    doc.add_paragraph(content)
                else:
                    doc.add_paragraph(section)
        
        # Save to memory
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # Return as downloadable file
        return StreamingResponse(
            file_stream, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=Proposal_{request.customer_name.replace(' ', '_')}.docx"}
        )
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Proposal generation failed: {str(e)}")

@app.get("/api/sales-insights")
def get_sales_insights():
    """Get sales performance insights and recommendations"""
    try:
        # Fetch recent sales data
        invoices_url = "https://bctest.dayliff.com:7048/BC160/ODataV4/Company('KENYA')/Sales_Invoice"
        invoices_response = requests.get(invoices_url, auth=get_auth())
        invoices_response.raise_for_status()
        invoices = invoices_response.json().get("value", [])
        
        # Example insights (in production, this would use actual analysis)
        insights = {
            "total_sales": len(invoices),
            "conversion_rate": "23.5%",  # This would come from actual calculations
            "top_performing_products": [
                {"name": "PKM60", "sales_count": 45, "revenue": 56700},
                {"name": "DDP 60", "sales_count": 32, "revenue": 41600},
                {"name": "GRUNDFOS CR", "sales_count": 28, "revenue": 89500}
            ],
            "common_objections": [
                {"objection": "Price too high", "frequency": "42%"},
                {"objection": "Delivery timeframe", "frequency": "28%"},
                {"objection": "Technical specifications", "frequency": "15%"}
            ],
            "recommendations": [
                "Focus on value proposition rather than price during negotiations",
                "Offer bundled solutions for higher conversion rates",
                "Schedule follow-ups within 3 days of initial contact"
            ]
        }
        
        return insights
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/recommend-products")
def recommend_products(customer_data: Customer):
    """Recommend products based on customer data and purchase history"""
    # In a production system, this would use an actual ML model
    # For demonstration, using a simple rule-based approach
    
    recommended_products = []
    
    # Check purchase history
    if customer_data.purchase_history:
        purchased_categories = set()
        for purchase in customer_data.purchase_history:
            product_category = purchase.get("category", "")
            if product_category:
                purchased_categories.add(product_category)
        
        # Recommend complementary products
        if "PUMPS" in purchased_categories:
            # Look up actual pump accessories
            try:
                url = "https://bctest.dayliff.com:7048/BC160/ODataV4/Company('KENYA')/Items"
                params = {"$filter": "contains(Description,'ACCESSORY')"}
                response = requests.get(url, params=params, auth=get_auth())
                accessories = response.json().get("value", [])
                
                for acc in accessories[:3]:  # Limit to 3 recommendations
                    recommended_products.append({
                        "id": acc.get("No"),
                        "name": acc.get("Description"),
                        "reason": "Complements your previous pump purchases"
                    })
            except Exception:
                # Fallback recommendations
                recommended_products.append({
                    "id": "PKM-ACC-01",
                    "name": "Pump Controller",
                    "reason": "Complements your previous pump purchases"
                })
    
    # Add industry-specific recommendations
    if customer_data.industry:
        industry_recommendations = {
            "agriculture": [{"id": "IRRIGATION-01", "name": "Irrigation System Kit", "reason": "Popular in your industry"}],
            "construction": [{"id": "DPUMP-01", "name": "Dewatering Pump Set", "reason": "Essential for construction projects"}],
            "residential": [{"id": "HWPUMP-01", "name": "Home Water System", "reason": "Perfect for residential applications"}]
        }
        
        industry = customer_data.industry.lower()
        if industry in industry_recommendations:
            recommended_products.extend(industry_recommendations[industry])
    
    # Add some generic recommendations if we don't have enough
    if len(recommended_products) < 3:
        generic_recommendations = [
            {"id": "PKM60", "name": "PKM60 Pump", "reason": "Bestselling product"},
            {"id": "DDP-60", "name": "DDP 60 Water Pump", "reason": "High performance and reliability"},
            {"id": "SOLAR-PUMP-01", "name": "Solar Pump System", "reason": "Energy-efficient solution"}
        ]
        
        # Add as many as needed to reach at least 3 recommendations
        for recommendation in generic_recommendations:
            if len(recommended_products) >= 3:
                break
            if not any(rec["id"] == recommendation["id"] for rec in recommended_products):
                recommended_products.append(recommendation)
    
    return {"recommendations": recommended_products}

@app.get("/api/optimal-follow-up")
def get_optimal_follow_up(customer_id: str):
    """Determine the optimal follow-up timing and strategy"""
    try:
        # Get customer interaction history
        url = BC_ERP_URL
        params = {"$filter": f"Contact_No eq '{customer_id}'"}
        
        response = requests.get(url, params=params, auth=get_auth())
        response.raise_for_status()
        
        interactions = response.json().get("value", [])
        
        # In production, this would use ML models to analyze patterns
        # For demonstration, using simple rules
        
        last_interaction = None
        interaction_count = len(interactions)
        
        if interactions:
            last_interaction = interactions[0]  # Assuming sorted by date
        
        current_date = datetime.now()
        
        if last_interaction:
            # Parse interaction date (adjust format as needed)
            try:
                last_date_str = last_interaction.get("Date", "")
                last_date = datetime.fromisoformat(last_date_str.replace("Z", "+00:00"))
            except:
                last_date = current_date - timedelta(days=random.randint(1, 14))
            
            days_since_last = (current_date - last_date).days
            
            # Logic for determining follow-up timing
            if days_since_last < 3:
                timing = "Wait 2 more days before following up"
            elif 3 <= days_since_last < 7:
                timing = "Follow up now - optimal timing"
            elif 7 <= days_since_last < 14:
                timing = "Follow up urgently - at risk of going cold"
            else:
                timing = "Re-engage with new value proposition"
        else:
            timing = "No previous interactions - initiate contact"
        
        # Determine method and message based on interaction count
        if interaction_count == 0:
            method = "Email"
            message = "Introduction and value proposition"
        elif interaction_count == 1:
            method = "Phone call"
            message = "Discuss specific needs and requirements"
        elif interaction_count == 2:
            method = "Email with proposal"
            message = "Share tailored solution based on previous discussions"
        elif interaction_count == 3:
            method = "In-person meeting"
            message = "Presentation and closing"
        else:
            method = "Personalized check-in"
            message = "Relationship maintenance and upsell opportunities"
        
        return {
            "timing": timing,
            "recommended_method": method,
            "suggested_message": message,
            "interaction_count": interaction_count,
            "days_since_last_contact": days_since_last if last_interaction else None
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)